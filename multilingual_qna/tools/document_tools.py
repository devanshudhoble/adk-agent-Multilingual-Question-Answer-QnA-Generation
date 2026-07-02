"""
Document Parsing Tools for ADK Agents
======================================
Provides tools for extracting text from PDF, DOCX, and TXT documents.
Uses ToolContext to persist extracted text in shared session state.
"""

import os
from PyPDF2 import PdfReader
from docx import Document


def parse_document(file_path: str) -> str:
    """Parse a document file and extract all text content.

    Supports PDF (.pdf), Word (.docx), and Plain Text (.txt) files.
    Use this tool when a user provides a file path and you need to
    extract the text content from the document.

    Args:
        file_path: The absolute path to the document file to parse.

    Returns:
        The extracted text content from the document, or an error message
        if the file cannot be read.
    """
    if not os.path.exists(file_path):
        return f"Error: File not found at '{file_path}'. Please verify the path."

    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".pdf":
            text = _extract_pdf(file_path)
        elif ext == ".docx":
            text = _extract_docx(file_path)
        elif ext == ".txt":
            text = _extract_txt(file_path)
        else:
            return (
                f"Error: Unsupported file format '{ext}'. "
                "Supported formats: .pdf, .docx, .txt"
            )

        if not text or len(text.strip()) < 50:
            return (
                "Error: Could not extract sufficient text from the document. "
                "The file may be empty, image-only, or corrupted."
            )

        # Truncate to stay within LLM context limits
        max_chars = 30000
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[... Document truncated at 30,000 characters ...]"

        char_count = len(text)
        return (
            f"DOCUMENT_METADATA:\n"
            f"- File: {filename}\n"
            f"- Format: {ext}\n"
            f"- Characters: {char_count:,}\n\n"
            f"DOCUMENT_CONTENT:\n{text}"
        )

    except Exception as e:
        return f"Error parsing '{filename}': {str(e)}"


def get_document_text(file_path: str) -> str:
    """Retrieve document text by parsing the file at the given path.

    This is a convenience tool that re-reads the document. Use this when
    you need the full document text content for QnA generation.

    Args:
        file_path: The absolute path to the document file.

    Returns:
        The raw text content extracted from the document.
    """
    if not os.path.exists(file_path):
        return f"Error: File not found at '{file_path}'."

    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(file_path)
        elif ext == ".docx":
            return _extract_docx(file_path)
        elif ext == ".txt":
            return _extract_txt(file_path)
        else:
            return f"Error: Unsupported format '{ext}'."
    except Exception as e:
        return f"Error: {str(e)}"


# ─── Internal Helper Functions ───────────────────────────────────────

def _extract_pdf(file_path: str) -> str:
    """Extract text from all pages of a PDF file."""
    reader = PdfReader(file_path)
    parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            parts.append(page_text.strip())
    return "\n\n".join(parts)


def _extract_docx(file_path: str) -> str:
    """Extract text from paragraphs and tables of a DOCX file."""
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def _extract_txt(file_path: str) -> str:
    """Extract text from a plain text file with encoding detection."""
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()
